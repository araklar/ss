"""
title: Word Table Generator Pipe
description: Generates a Word document with a table from LLM output.
author: Prometheus
version: 0.1.0
license: MIT
"""

from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
from utils.misc import get_last_assistant_message


class Pipeline:
    def __init__(self):
        self.name = "Word Table Generator"

    def pipe(self, user_message: str, model_id: str, messages: list, body: dict) -> str:
        llm_output = get_last_assistant_message(messages)

        if "generate_word_table" in user_message:
            # LLM çıktısını tabloya dönüştür
            table_data = self.extract_table_data(llm_output)

            # Word belgesi oluştur
            doc = Document()
            doc.add_heading("LLM Çıktısı Tablosu", 0)

            # Tabloyu oluştur
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(table_data[0]):
                hdr_cells[i].text = header

            # Verileri tabloya ekle
            for row in table_data[1:]:
                cells = table.add_row().cells
                for i, cell_data in enumerate(row):
                    cells[i].text = cell_data

            # Tablo stilini ve yazı boyutunu ayarla
            table.style = "Table Grid"
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            # Belgeyi BytesIO'ya kaydet
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            # Open WebUI'ye dosya olarak döndür
            return {
                "type": "file",
                "data": {
                    "filename": "llm_table.docx",
                    "content": doc_bytes.read(),
                    "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                },
            }
        else:
            return llm_output

    def extract_table_data(self, text: str) -> list:
        """
        LLM çıktısından tablo verilerini çıkarır.
        Bu fonksiyonu LLM çıktınızın formatına göre özelleştirmeniz gerekebilir.
        """
        # Varsayılan olarak, her satırı virgülle ayrılmış değerler olarak kabul ediyoruz.
        table_data = [line.split(",") for line in text.splitlines()]
        return table_data
